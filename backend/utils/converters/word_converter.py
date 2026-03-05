from .base import BaseConverter
from docx import Document
import aiofiles
import os


class WordConverter(BaseConverter):
    """Word 文档转换器"""

    @classmethod
    def supports(cls, filename: str) -> bool:
        """判断是否支持该文件类型"""
        return filename.lower().endswith((".docx", ".doc"))

    @classmethod
    async def convert(cls, file_path: str) -> tuple[str, str]:
        """
        转换 Word 文档为 Markdown
        :return: (markdown_content, title)
        """
        # 使用同步方式读取 docx 文件
        doc = Document(file_path)

        # 提取标题
        title = "未命名文档"
        if doc.paragraphs:
            first_para = doc.paragraphs[0].text.strip()
            if first_para:
                title = first_para

        # 转换为 Markdown
        markdown_lines = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # 根据样式格式化
            if para.style.name.startswith("Heading"):
                level = para.style.name.replace("Heading ", "")
                markdown_lines.append(f"{'#' * int(level)} {text}")
            else:
                markdown_lines.append(text)

        content = "\n\n".join(markdown_lines)
        return content, title
